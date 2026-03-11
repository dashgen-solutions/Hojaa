"""
Document Builder API — block-based proposals, SOWs, and contracts.

REST endpoints for document CRUD, content auto-save, versioning, pricing,
sharing/recipients, templates, variable resolution, and PDF export.
"""
from __future__ import annotations

import io
import secrets
from datetime import datetime, date
from typing import List, Optional
from uuid import UUID

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from sqlalchemy import func, desc
from sqlalchemy.orm import Session, joinedload

from app.core.auth import get_current_user
from app.core.logger import get_logger
from app.db.session import get_db
from app.models.database import (
    User, Organization, OrgRole,
    Session as SessionModel, Node, Card, TeamMember,
    Document, DocumentVersion, DocumentTemplate,
    DocumentRecipient, PricingLineItem, DocumentStatus,
    DocumentChatMessage,
)

logger = get_logger(__name__)
router = APIRouter(prefix="/documents", tags=["documents"])


# -- Pydantic Schemas --------------------------------------------------------

class CreateDocumentRequest(BaseModel):
    title: Optional[str] = None
    template_id: Optional[str] = None


class UpdateDocumentRequest(BaseModel):
    title: Optional[str] = None
    status: Optional[str] = None


class SaveContentRequest(BaseModel):
    content: list


class CreateVersionRequest(BaseModel):
    change_summary: Optional[str] = None


class RenameVersionRequest(BaseModel):
    change_summary: str = Field(..., min_length=1, max_length=500)


class AddLineItemRequest(BaseModel):
    name: str = Field(..., min_length=1, max_length=500)
    description: Optional[str] = None
    quantity: float = 1.0
    unit_price: float = 0.0
    discount_percent: float = 0.0
    tax_percent: float = 0.0
    is_optional: bool = False
    order_index: int = 0


class UpdateLineItemRequest(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    quantity: Optional[float] = None
    unit_price: Optional[float] = None
    discount_percent: Optional[float] = None
    tax_percent: Optional[float] = None
    is_optional: Optional[bool] = None
    is_selected: Optional[bool] = None
    order_index: Optional[int] = None


class PricingFromCardsRequest(BaseModel):
    hourly_rate: float = Field(..., gt=0)


class RecipientInput(BaseModel):
    name: str = Field(..., min_length=1, max_length=255)
    email: str = Field(..., min_length=1, max_length=255)
    role: str = "viewer"


class ShareDocumentRequest(BaseModel):
    recipients: Optional[List[RecipientInput]] = None


class CreateTemplateRequest(BaseModel):
    name: str = Field(..., min_length=1, max_length=255)
    description: Optional[str] = None
    category: Optional[str] = None


# -- Helpers ------------------------------------------------------------------

def _verify_session_access(db: Session, session_id: str, user: User) -> SessionModel:
    """Load a session and verify the user has access to it."""
    session = db.query(SessionModel).filter(SessionModel.id == session_id).first()
    if not session:
        raise HTTPException(404, "Session not found")
    # Owner always has access
    if session.user_id == user.id:
        return session
    # Same organization
    if session.organization_id and session.organization_id == user.organization_id:
        return session
    raise HTTPException(403, "You do not have access to this session")


def _verify_document_access(db: Session, document_id: str, user: User) -> Document:
    """Load a document and verify the user has access to it."""
    doc = (
        db.query(Document)
        .filter(Document.id == document_id)
        .first()
    )
    if not doc:
        raise HTTPException(404, "Document not found")
    # Creator always has access
    if doc.created_by == user.id:
        return doc
    # Same organization
    if doc.organization_id and doc.organization_id == user.organization_id:
        return doc
    raise HTTPException(403, "You do not have access to this document")


def _is_admin_or_creator(user: User, doc: Document) -> bool:
    """Check if user is the document creator or an organization admin/owner."""
    if doc.created_by == user.id:
        return True
    if user.org_role in (OrgRole.ADMIN, OrgRole.OWNER):
        return True
    return False


def _serialize_document_summary(doc: Document) -> dict:
    return {
        "id": str(doc.id),
        "title": doc.title,
        "status": doc.status.value if doc.status else "draft",
        "creator_name": doc.creator.username if doc.creator else "Unknown",
        "recipients_count": len(doc.recipients) if doc.recipients else 0,
        "updated_at": doc.updated_at.isoformat() if doc.updated_at else None,
        "created_at": doc.created_at.isoformat() if doc.created_at else None,
    }


def _serialize_document_full(doc: Document) -> dict:
    return {
        "id": str(doc.id),
        "session_id": str(doc.session_id),
        "organization_id": str(doc.organization_id),
        "created_by": str(doc.created_by) if doc.created_by else None,
        "creator_name": doc.creator.username if doc.creator else "Unknown",
        "title": doc.title,
        "status": doc.status.value if doc.status else "draft",
        "template_id": str(doc.template_id) if doc.template_id else None,
        "content": doc.content or [],
        "share_token": doc.share_token,
        "expires_at": doc.expires_at.isoformat() if doc.expires_at else None,
        "sent_at": doc.sent_at.isoformat() if doc.sent_at else None,
        "viewed_at": doc.viewed_at.isoformat() if doc.viewed_at else None,
        "completed_at": doc.completed_at.isoformat() if doc.completed_at else None,
        "created_at": doc.created_at.isoformat() if doc.created_at else None,
        "updated_at": doc.updated_at.isoformat() if doc.updated_at else None,
        "recipients": [
            {
                "id": str(r.id),
                "name": r.name,
                "email": r.email,
                "role": r.role,
                "sent_at": r.sent_at.isoformat() if r.sent_at else None,
                "viewed_at": r.viewed_at.isoformat() if r.viewed_at else None,
                "completed_at": r.completed_at.isoformat() if r.completed_at else None,
            }
            for r in (doc.recipients or [])
        ],
        "pricing_items": [
            _serialize_line_item(item)
            for item in sorted((doc.pricing_items or []), key=lambda i: i.order_index)
        ],
    }


def _serialize_line_item(item: PricingLineItem) -> dict:
    line_total = item.quantity * item.unit_price * (1 - item.discount_percent / 100) * (1 + item.tax_percent / 100)
    return {
        "id": str(item.id),
        "document_id": str(item.document_id),
        "name": item.name,
        "description": item.description,
        "quantity": item.quantity,
        "unit_price": item.unit_price,
        "discount_percent": item.discount_percent,
        "tax_percent": item.tax_percent,
        "is_optional": item.is_optional,
        "is_selected": item.is_selected,
        "order_index": item.order_index,
        "card_id": str(item.card_id) if item.card_id else None,
        "line_total": round(line_total, 2),
        "created_at": item.created_at.isoformat() if item.created_at else None,
        "updated_at": item.updated_at.isoformat() if item.updated_at else None,
    }


def _resolve_variables(db: Session, doc: Document) -> dict:
    """Resolve template variables to current project data."""
    session = db.query(SessionModel).filter(SessionModel.id == doc.session_id).first()
    organization = db.query(Organization).filter(Organization.id == doc.organization_id).first() if doc.organization_id else None

    team_members = (
        db.query(TeamMember)
        .filter(TeamMember.session_id == doc.session_id)
        .all()
    )

    variables = {
        "project.name": session.document_filename if session and session.document_filename else "Untitled Project",
        "project.status": session.status.value if session and session.status else "",
        "project.created_at": session.created_at.isoformat() if session and session.created_at else "",
        "org.name": organization.name if organization else "",
        "org.industry": organization.industry if organization else "",
        "org.website": organization.website if organization else "",
        "date.today": date.today().isoformat(),
        "date.year": str(date.today().year),
        "document.title": doc.title or "",
        "document.status": doc.status.value if doc.status else "",
        "team.members": ", ".join(m.name for m in team_members),
        "team.count": str(len(team_members)),
    }

    # Add individual team member entries
    for idx, member in enumerate(team_members):
        variables[f"team.member.{idx}.name"] = member.name
        variables[f"team.member.{idx}.email"] = member.email or ""
        variables[f"team.member.{idx}.role"] = member.role or ""

    return variables


# -- Document CRUD ------------------------------------------------------------

@router.get("/session/{session_id}")
def list_documents(
    session_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """List documents for a project session."""
    _verify_session_access(db, session_id, current_user)

    docs = (
        db.query(Document)
        .filter(Document.session_id == session_id)
        .options(
            joinedload(Document.creator),
            joinedload(Document.recipients),
        )
        .order_by(desc(Document.updated_at))
        .all()
    )

    return [_serialize_document_summary(d) for d in docs]


@router.post("/session/{session_id}", status_code=201)
def create_document(
    session_id: str,
    body: CreateDocumentRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Create a new document (blank or from a template)."""
    session = _verify_session_access(db, session_id, current_user)

    title = body.title or "Untitled Document"
    content: list = []
    template_id = None

    # If creating from a template, copy its content
    if body.template_id:
        template = db.query(DocumentTemplate).filter(
            DocumentTemplate.id == body.template_id,
            DocumentTemplate.is_active == True,
        ).first()
        if not template:
            raise HTTPException(404, "Template not found")
        content = list(template.content) if template.content else []
        template_id = template.id
        if not body.title and template.name:
            title = template.name

    doc = Document(
        session_id=UUID(session_id),
        organization_id=session.organization_id or current_user.organization_id,
        created_by=current_user.id,
        title=title,
        status=DocumentStatus.DRAFT,
        template_id=template_id,
        content=content,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    # If created from template, resolve variables in the content
    if body.template_id and content:
        variables = _resolve_variables(db, doc)
        resolved = _apply_variables_to_content(content, variables)
        doc.content = resolved
        db.commit()
        db.refresh(doc)

    return {"id": str(doc.id), "title": doc.title, "status": doc.status.value}


def _apply_variables_to_content(content: list, variables: dict) -> list:
    """Recursively replace {{variable}} placeholders in content blocks."""
    import json as json_lib
    serialized = json_lib.dumps(content)
    for key, value in variables.items():
        serialized = serialized.replace("{{" + key + "}}", str(value))
    return json_lib.loads(serialized)


def _generate_version_summary(db: Session, doc: Document, prev_version_number: int) -> str:
    """Generate a descriptive summary for a version snapshot by analysing the content.
    
    Extracts section headings from the current content and optionally compares 
    with the previous version to describe what changed.
    """
    # Extract section headings from current content
    headings = []
    block_count = 0
    if doc.content and isinstance(doc.content, list):
        block_count = len(doc.content)
        for block in doc.content:
            if isinstance(block, dict) and block.get("type") == "heading":
                text_parts = []
                for item in (block.get("content") or []):
                    if isinstance(item, dict):
                        text_parts.append(item.get("text", ""))
                heading_text = "".join(text_parts).strip()
                if heading_text:
                    headings.append(heading_text)
    
    if not headings and block_count == 0:
        return "Empty document snapshot"
    
    # Try to compare with previous version
    if prev_version_number > 0:
        prev = (
            db.query(DocumentVersion)
            .filter(
                DocumentVersion.document_id == doc.id,
                DocumentVersion.version_number == prev_version_number,
            )
            .first()
        )
        if prev and prev.content:
            prev_headings = set()
            for block in prev.content:
                if isinstance(block, dict) and block.get("type") == "heading":
                    text_parts = []
                    for item in (block.get("content") or []):
                        if isinstance(item, dict):
                            text_parts.append(item.get("text", ""))
                    h = "".join(text_parts).strip()
                    if h:
                        prev_headings.add(h)
            
            new_sections = [h for h in headings if h not in prev_headings]
            prev_block_count = len(prev.content) if prev.content else 0
            diff = block_count - prev_block_count
            
            parts = []
            if new_sections:
                parts.append(f"Added: {', '.join(new_sections[:3])}")
                if len(new_sections) > 3:
                    parts.append(f"(+{len(new_sections) - 3} more sections)")
            if diff > 0:
                parts.append(f"+{diff} blocks")
            elif diff < 0:
                parts.append(f"{diff} blocks")
            
            if parts:
                return "; ".join(parts)
    
    # First version or no previous — describe content
    if headings:
        sections_str = ", ".join(headings[:4])
        if len(headings) > 4:
            sections_str += f" (+{len(headings) - 4} more)"
        return f"Sections: {sections_str}"
    
    return f"Document snapshot ({block_count} blocks)"


def _extract_section_headings(content) -> list:
    """Extract heading text from BlockNote content blocks for version preview."""
    headings = []
    if not content or not isinstance(content, list):
        return headings
    for block in content:
        if isinstance(block, dict) and block.get("type") == "heading":
            text_parts = []
            for item in (block.get("content") or []):
                if isinstance(item, dict):
                    text_parts.append(item.get("text", ""))
            heading_text = "".join(text_parts).strip()
            if heading_text:
                headings.append(heading_text)
    return headings


# -- Templates (MUST be registered before /{document_id} to avoid route conflict) --

@router.get("/templates")
def list_templates(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """List available templates (system templates + user's org templates)."""
    query = db.query(DocumentTemplate).filter(
        DocumentTemplate.is_active == True,
    )

    if current_user.organization_id:
        query = query.filter(
            (DocumentTemplate.is_system == True)
            | (DocumentTemplate.organization_id == current_user.organization_id)
        )
    else:
        query = query.filter(DocumentTemplate.is_system == True)

    templates = query.order_by(DocumentTemplate.category, DocumentTemplate.name).all()

    return [
        {
            "id": str(t.id),
            "name": t.name,
            "description": t.description,
            "category": t.category,
            "thumbnail_url": t.thumbnail_url,
            "is_system": t.is_system,
            "variables": t.variables or [],
            "created_at": t.created_at.isoformat() if t.created_at else None,
        }
        for t in templates
    ]


@router.post("/templates", status_code=201)
def create_template(
    body: CreateTemplateRequest,
    document_id: str = Query(..., description="Source document ID to save as template"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Save an existing document as a reusable template."""
    doc = _verify_document_access(db, document_id, current_user)

    template = DocumentTemplate(
        organization_id=current_user.organization_id,
        created_by=current_user.id,
        name=body.name,
        description=body.description,
        category=body.category,
        content=list(doc.content) if doc.content else [],
        variables=[],
        is_system=False,
        is_active=True,
    )
    db.add(template)
    db.commit()
    db.refresh(template)

    return {
        "id": str(template.id),
        "name": template.name,
        "description": template.description,
        "category": template.category,
        "created_at": template.created_at.isoformat(),
    }


@router.delete("/templates/{template_id}", status_code=204)
def delete_template(
    template_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Delete an organization template (not system templates)."""
    template = db.query(DocumentTemplate).filter(DocumentTemplate.id == template_id).first()
    if not template:
        raise HTTPException(404, "Template not found")

    if template.is_system:
        raise HTTPException(403, "Cannot delete system templates")

    # Only org members or the creator can delete
    if template.organization_id != current_user.organization_id and template.created_by != current_user.id:
        raise HTTPException(403, "You do not have permission to delete this template")

    db.delete(template)
    db.commit()


# -- Document by ID -----------------------------------------------------------

@router.get("/{document_id}")
def get_document(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get full document with content, pricing items, and recipients."""
    doc = (
        db.query(Document)
        .filter(Document.id == document_id)
        .options(
            joinedload(Document.creator),
            joinedload(Document.recipients),
            joinedload(Document.pricing_items),
        )
        .first()
    )
    if not doc:
        raise HTTPException(404, "Document not found")

    # Access check
    if doc.created_by != current_user.id:
        if not (doc.organization_id and doc.organization_id == current_user.organization_id):
            raise HTTPException(403, "You do not have access to this document")

    return _serialize_document_full(doc)


@router.patch("/{document_id}")
def update_document(
    document_id: str,
    body: UpdateDocumentRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Update document title or status."""
    doc = _verify_document_access(db, document_id, current_user)

    if body.title is not None:
        doc.title = body.title
    if body.status is not None:
        try:
            new_status = DocumentStatus(body.status)
        except ValueError:
            raise HTTPException(400, f"Invalid status: {body.status}")
        doc.status = new_status
        # Update lifecycle timestamps
        if new_status == DocumentStatus.COMPLETED:
            doc.completed_at = datetime.utcnow()

    db.commit()
    db.refresh(doc)

    return {
        "id": str(doc.id),
        "title": doc.title,
        "status": doc.status.value,
        "updated_at": doc.updated_at.isoformat(),
    }


@router.delete("/{document_id}", status_code=204)
def delete_document(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Delete a document (creator or admin only)."""
    doc = _verify_document_access(db, document_id, current_user)

    if not _is_admin_or_creator(current_user, doc):
        raise HTTPException(403, "Only the document creator or an admin can delete it")

    db.delete(doc)
    db.commit()


@router.post("/{document_id}/duplicate", status_code=201)
def duplicate_document(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Clone an existing document."""
    doc = _verify_document_access(db, document_id, current_user)

    clone = Document(
        session_id=doc.session_id,
        organization_id=doc.organization_id,
        created_by=current_user.id,
        title=f"{doc.title} (Copy)",
        status=DocumentStatus.DRAFT,
        template_id=doc.template_id,
        content=list(doc.content) if doc.content else [],
    )
    db.add(clone)
    db.flush()

    # Duplicate pricing items
    existing_items = (
        db.query(PricingLineItem)
        .filter(PricingLineItem.document_id == doc.id)
        .order_by(PricingLineItem.order_index)
        .all()
    )
    for item in existing_items:
        db.add(PricingLineItem(
            document_id=clone.id,
            name=item.name,
            description=item.description,
            quantity=item.quantity,
            unit_price=item.unit_price,
            discount_percent=item.discount_percent,
            tax_percent=item.tax_percent,
            is_optional=item.is_optional,
            is_selected=item.is_selected,
            order_index=item.order_index,
            card_id=item.card_id,
        ))

    db.commit()
    db.refresh(clone)

    return {"id": str(clone.id), "title": clone.title, "status": clone.status.value}


# -- Content & Auto-save -----------------------------------------------------

@router.put("/{document_id}/content")
def save_content(
    document_id: str,
    body: SaveContentRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Auto-save document content (BlockNote JSON blocks)."""
    doc = _verify_document_access(db, document_id, current_user)

    doc.content = body.content
    db.commit()
    db.refresh(doc)

    return {"status": "saved", "updated_at": doc.updated_at.isoformat()}


@router.get("/{document_id}/versions")
def list_versions(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """List version snapshots for a document."""
    _verify_document_access(db, document_id, current_user)

    versions = (
        db.query(DocumentVersion)
        .filter(DocumentVersion.document_id == document_id)
        .options(joinedload(DocumentVersion.author))
        .order_by(desc(DocumentVersion.version_number))
        .all()
    )

    return [
        {
            "id": str(v.id),
            "version_number": v.version_number,
            "change_summary": v.change_summary,
            "changed_by": str(v.changed_by) if v.changed_by else None,
            "author_name": v.author.username if v.author else "Unknown",
            "content_preview": _extract_section_headings(v.content),
            "created_at": v.created_at.isoformat(),
        }
        for v in versions
    ]


@router.patch("/{document_id}/versions/{version_id}")
def rename_version(
    document_id: str,
    version_id: str,
    body: RenameVersionRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Rename a version (update its change_summary)."""
    _verify_document_access(db, document_id, current_user)

    version = (
        db.query(DocumentVersion)
        .filter(
            DocumentVersion.id == version_id,
            DocumentVersion.document_id == document_id,
        )
        .options(joinedload(DocumentVersion.author))
        .first()
    )
    if not version:
        raise HTTPException(404, "Version not found")

    version.change_summary = body.change_summary
    db.commit()
    db.refresh(version)

    return {
        "id": str(version.id),
        "version_number": version.version_number,
        "change_summary": version.change_summary,
        "changed_by": str(version.changed_by) if version.changed_by else None,
        "author_name": version.author.username if version.author else "Unknown",
        "content_preview": _extract_section_headings(version.content),
        "created_at": version.created_at.isoformat(),
    }


@router.post("/{document_id}/versions", status_code=201)
def create_version(
    document_id: str,
    body: CreateVersionRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Create a named version snapshot of the current content.
    
    If no change_summary is provided, auto-generates one by analysing
    the document content (compares with previous version if one exists).
    """
    doc = _verify_document_access(db, document_id, current_user)

    # Determine next version number
    max_version = (
        db.query(func.max(DocumentVersion.version_number))
        .filter(DocumentVersion.document_id == doc.id)
        .scalar()
    ) or 0

    # Auto-generate a descriptive summary when none is provided
    summary = body.change_summary
    if not summary or summary.strip() in ("", "Manual snapshot"):
        summary = _generate_version_summary(db, doc, max_version)

    version = DocumentVersion(
        document_id=doc.id,
        version_number=max_version + 1,
        content=list(doc.content) if doc.content else [],
        changed_by=current_user.id,
        change_summary=summary,
    )
    db.add(version)
    db.commit()
    db.refresh(version)

    return {
        "id": str(version.id),
        "version_number": version.version_number,
        "change_summary": version.change_summary,
        "author_name": current_user.username if current_user else "Unknown",
        "created_at": version.created_at.isoformat(),
    }


@router.get("/{document_id}/versions/{version_id}")
def get_version(
    document_id: str,
    version_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get a specific version snapshot with its content."""
    _verify_document_access(db, document_id, current_user)

    version = (
        db.query(DocumentVersion)
        .filter(
            DocumentVersion.id == version_id,
            DocumentVersion.document_id == document_id,
        )
        .options(joinedload(DocumentVersion.author))
        .first()
    )
    if not version:
        raise HTTPException(404, "Version not found")

    return {
        "id": str(version.id),
        "document_id": str(version.document_id),
        "version_number": version.version_number,
        "content": version.content,
        "change_summary": version.change_summary,
        "changed_by": str(version.changed_by) if version.changed_by else None,
        "author_name": version.author.username if version.author else "Unknown",
        "created_at": version.created_at.isoformat(),
    }


@router.post("/{document_id}/versions/{version_id}/restore")
def restore_version(
    document_id: str,
    version_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Restore a document to a previous version.

    Replaces the document content with the selected version's content.
    """
    doc = _verify_document_access(db, document_id, current_user)

    version = (
        db.query(DocumentVersion)
        .filter(
            DocumentVersion.id == version_id,
            DocumentVersion.document_id == document_id,
        )
        .first()
    )
    if not version:
        raise HTTPException(404, "Version not found")

    # Replace current content with the version's content
    doc.content = list(version.content) if version.content else []
    db.commit()
    db.refresh(doc)

    return {
        "status": "restored",
        "restored_version": version.version_number,
        "content": doc.content,
    }


# -- Variables & Scope Integration -------------------------------------------

@router.get("/{document_id}/variables")
def get_variables(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Resolve template variables to current project data."""
    doc = _verify_document_access(db, document_id, current_user)
    variables = _resolve_variables(db, doc)
    return {"variables": variables}


@router.get("/session/{session_id}/scope-data")
def get_scope_data(
    session_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get scope tree nodes and planning cards for reference insertion."""
    _verify_session_access(db, session_id, current_user)

    # Scope tree nodes
    nodes = (
        db.query(Node)
        .filter(Node.session_id == session_id)
        .order_by(Node.depth, Node.order_index)
        .all()
    )

    # Planning cards
    cards = (
        db.query(Card)
        .filter(Card.session_id == session_id)
        .order_by(Card.created_at)
        .all()
    )

    return {
        "nodes": [
            {
                "id": str(n.id),
                "parent_id": str(n.parent_id) if n.parent_id else None,
                "question": n.question,
                "answer": n.answer,
                "node_type": n.node_type.value if n.node_type else None,
                "status": n.status.value if n.status else None,
                "depth": n.depth,
                "order_index": n.order_index,
            }
            for n in nodes
        ],
        "cards": [
            {
                "id": str(c.id),
                "node_id": str(c.node_id) if c.node_id else None,
                "title": c.title or (c.node.question if c.node else "Untitled"),
                "description": c.description or (c.node.answer if c.node else None),
                "status": c.status.value if c.status else None,
                "priority": c.priority.value if c.priority else None,
                "estimated_hours": c.estimated_hours,
            }
            for c in cards
        ],
    }


# -- Pricing ------------------------------------------------------------------

@router.get("/{document_id}/pricing")
def get_pricing(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get pricing line items with calculated totals."""
    _verify_document_access(db, document_id, current_user)

    items = (
        db.query(PricingLineItem)
        .filter(PricingLineItem.document_id == document_id)
        .order_by(PricingLineItem.order_index)
        .all()
    )

    serialized = [_serialize_line_item(i) for i in items]

    # Calculate totals
    subtotal = sum(i["line_total"] for i in serialized if not i["is_optional"] or i["is_selected"])
    optional_total = sum(i["line_total"] for i in serialized if i["is_optional"] and not i["is_selected"])

    return {
        "items": serialized,
        "subtotal": round(subtotal, 2),
        "optional_total": round(optional_total, 2),
        "total": round(subtotal, 2),
    }


@router.post("/{document_id}/pricing", status_code=201)
def add_line_item(
    document_id: str,
    body: AddLineItemRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Add a pricing line item to a document."""
    _verify_document_access(db, document_id, current_user)

    item = PricingLineItem(
        document_id=UUID(document_id),
        name=body.name,
        description=body.description,
        quantity=body.quantity,
        unit_price=body.unit_price,
        discount_percent=body.discount_percent,
        tax_percent=body.tax_percent,
        is_optional=body.is_optional,
        order_index=body.order_index,
    )
    db.add(item)
    db.commit()
    db.refresh(item)

    return _serialize_line_item(item)


@router.patch("/pricing/{item_id}")
def update_line_item(
    item_id: str,
    body: UpdateLineItemRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Update a pricing line item."""
    item = db.query(PricingLineItem).filter(PricingLineItem.id == item_id).first()
    if not item:
        raise HTTPException(404, "Line item not found")

    # Verify access via parent document
    _verify_document_access(db, str(item.document_id), current_user)

    if body.name is not None:
        item.name = body.name
    if body.description is not None:
        item.description = body.description
    if body.quantity is not None:
        item.quantity = body.quantity
    if body.unit_price is not None:
        item.unit_price = body.unit_price
    if body.discount_percent is not None:
        item.discount_percent = body.discount_percent
    if body.tax_percent is not None:
        item.tax_percent = body.tax_percent
    if body.is_optional is not None:
        item.is_optional = body.is_optional
    if body.is_selected is not None:
        item.is_selected = body.is_selected
    if body.order_index is not None:
        item.order_index = body.order_index

    db.commit()
    db.refresh(item)

    return _serialize_line_item(item)


@router.delete("/pricing/{item_id}", status_code=204)
def delete_line_item(
    item_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Remove a pricing line item."""
    item = db.query(PricingLineItem).filter(PricingLineItem.id == item_id).first()
    if not item:
        raise HTTPException(404, "Line item not found")

    _verify_document_access(db, str(item.document_id), current_user)

    db.delete(item)
    db.commit()


@router.post("/{document_id}/pricing/from-cards", status_code=201)
def pricing_from_cards(
    document_id: str,
    body: PricingFromCardsRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Generate pricing line items from planning cards (estimated_hours as quantity)."""
    doc = _verify_document_access(db, document_id, current_user)

    cards = (
        db.query(Card)
        .filter(
            Card.session_id == doc.session_id,
            Card.estimated_hours != None,
            Card.estimated_hours > 0,
        )
        .options(joinedload(Card.node))
        .all()
    )

    if not cards:
        raise HTTPException(400, "No planning cards with estimated hours found")

    created = []
    for idx, card in enumerate(cards):
        card_title = card.title or (card.node.question if card.node else "Untitled Task")
        card_desc = card.description or (card.node.answer if card.node else None)

        item = PricingLineItem(
            document_id=doc.id,
            name=card_title,
            description=card_desc,
            quantity=card.estimated_hours,
            unit_price=body.hourly_rate,
            discount_percent=0.0,
            tax_percent=0.0,
            is_optional=False,
            is_selected=True,
            order_index=idx,
            card_id=card.id,
        )
        db.add(item)
        created.append(item)

    db.commit()

    # Refresh and serialize
    for item in created:
        db.refresh(item)

    return {
        "count": len(created),
        "items": [_serialize_line_item(i) for i in created],
    }


# -- Sharing & Recipients ----------------------------------------------------

@router.post("/{document_id}/share")
def share_document(
    document_id: str,
    body: ShareDocumentRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Generate a share link and optionally add recipients."""
    doc = _verify_document_access(db, document_id, current_user)

    # Generate share token if not already set
    if not doc.share_token:
        doc.share_token = secrets.token_urlsafe(32)

    added_recipients = []
    if body.recipients:
        for r in body.recipients:
            recipient = DocumentRecipient(
                document_id=doc.id,
                name=r.name,
                email=r.email,
                role=r.role,
                access_token=secrets.token_urlsafe(32),
            )
            db.add(recipient)
            added_recipients.append(recipient)

    db.commit()
    db.refresh(doc)

    # Refresh recipients for serialization
    for r in added_recipients:
        db.refresh(r)

    # Return ALL recipients (not just newly added) so the frontend stays in sync
    all_recipients = (
        db.query(DocumentRecipient)
        .filter(DocumentRecipient.document_id == doc.id)
        .order_by(DocumentRecipient.created_at)
        .all()
    )

    return {
        "share_token": doc.share_token,
        "recipients": [
            {
                "id": str(r.id),
                "name": r.name,
                "email": r.email,
                "role": r.role,
                "access_token": r.access_token,
                "sent_at": r.sent_at.isoformat() if r.sent_at else None,
                "viewed_at": r.viewed_at.isoformat() if r.viewed_at else None,
                "completed_at": r.completed_at.isoformat() if r.completed_at else None,
                "created_at": r.created_at.isoformat() if r.created_at else None,
            }
            for r in all_recipients
        ],
    }


@router.delete("/{document_id}/recipients/{recipient_id}")
def remove_recipient(
    document_id: str,
    recipient_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Remove a recipient from a document."""
    _verify_document_access(db, document_id, current_user)

    recipient = (
        db.query(DocumentRecipient)
        .filter(
            DocumentRecipient.id == recipient_id,
            DocumentRecipient.document_id == document_id,
        )
        .first()
    )
    if not recipient:
        raise HTTPException(status_code=404, detail="Recipient not found")

    db.delete(recipient)
    db.commit()
    return {"success": True}


@router.get("/{document_id}/recipients")
def list_recipients(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """List recipients with their status."""
    _verify_document_access(db, document_id, current_user)

    recipients = (
        db.query(DocumentRecipient)
        .filter(DocumentRecipient.document_id == document_id)
        .order_by(DocumentRecipient.created_at)
        .all()
    )

    return [
        {
            "id": str(r.id),
            "name": r.name,
            "email": r.email,
            "role": r.role,
            "sent_at": r.sent_at.isoformat() if r.sent_at else None,
            "viewed_at": r.viewed_at.isoformat() if r.viewed_at else None,
            "completed_at": r.completed_at.isoformat() if r.completed_at else None,
            "access_token": r.access_token,
            "created_at": r.created_at.isoformat() if r.created_at else None,
        }
        for r in recipients
    ]


@router.post("/{document_id}/send")
def send_document(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Mark document as sent and update sent_at timestamps on document and recipients."""
    doc = _verify_document_access(db, document_id, current_user)

    now = datetime.utcnow()

    doc.status = DocumentStatus.SENT
    doc.sent_at = now

    # Update all recipients
    recipients = (
        db.query(DocumentRecipient)
        .filter(DocumentRecipient.document_id == doc.id)
        .all()
    )
    for r in recipients:
        if not r.sent_at:
            r.sent_at = now

    db.commit()
    db.refresh(doc)

    return {
        "id": str(doc.id),
        "status": doc.status.value,
        "sent_at": doc.sent_at.isoformat(),
        "recipients_sent": len(recipients),
    }


@router.get("/view/{access_token}")
def view_document_public(
    access_token: str,
    db: Session = Depends(get_db),
):
    """PUBLIC endpoint — view a document via share_token or recipient access_token.

    Accepts either the document-level share_token (from the copy-link flow)
    or a per-recipient access_token.
    """
    # Try per-recipient access_token first
    recipient = (
        db.query(DocumentRecipient)
        .filter(DocumentRecipient.access_token == access_token)
        .first()
    )

    if recipient:
        doc = (
            db.query(Document)
            .filter(Document.id == recipient.document_id)
            .options(joinedload(Document.creator), joinedload(Document.pricing_items))
            .first()
        )
    else:
        # Fallback: try document-level share_token
        doc = (
            db.query(Document)
            .filter(Document.share_token == access_token)
            .options(joinedload(Document.creator), joinedload(Document.pricing_items))
            .first()
        )

    if not doc:
        raise HTTPException(404, "Invalid or expired access link")

    # Check expiry
    if doc.expires_at and doc.expires_at < datetime.utcnow():
        raise HTTPException(410, "This document link has expired")

    # Update viewed_at timestamps
    now = datetime.utcnow()
    if recipient and not recipient.viewed_at:
        recipient.viewed_at = now
    if not doc.viewed_at:
        doc.viewed_at = now
        if doc.status == DocumentStatus.SENT:
            doc.status = DocumentStatus.VIEWED

    db.commit()
    db.refresh(doc)

    pricing_items = sorted((doc.pricing_items or []), key=lambda i: i.order_index)

    return {
        "title": doc.title,
        "status": doc.status.value,
        "content": doc.content or [],
        "creator_name": doc.creator.username if doc.creator else "Unknown",
        "sent_at": doc.sent_at.isoformat() if doc.sent_at else None,
        "recipient": {
            "name": recipient.name if recipient else "Link visitor",
            "email": recipient.email if recipient else "",
            "role": recipient.role if recipient else "viewer",
        },
        "pricing_items": [_serialize_line_item(i) for i in pricing_items],
    }


# -- Export (PDF & DOCX) ------------------------------------------------------


class ExportFormatQuery:
    """Allow ?format=pdf or ?format=docx query param."""
    pass


@router.get("/{document_id}/docx")
def export_document_docx(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Render document to DOCX (Word) and return as a streaming download."""
    doc = _verify_document_access(db, document_id, current_user)
    content_blocks = doc.content or []

    try:
        from docx import Document as WordDocument
        from docx.shared import Pt, Inches, RGBColor, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        raise HTTPException(
            500,
            "DOCX generation requires the python-docx package. Install with: pip install python-docx",
        )

    word = WordDocument()

    # ── Configure default styles to match preview ──
    style_normal = word.styles["Normal"]
    style_normal.font.name = "Inter"
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = RGBColor(55, 65, 81)  # text-neutral-700
    style_normal.paragraph_format.space_after = Pt(4)
    style_normal.paragraph_format.line_spacing = 1.5

    # Heading styles
    for level_num, (size, color_hex) in {
        0: (24, "111827"),
        1: (20, "111827"),
        2: (16, "1F2937"),
        3: (13, "374151"),
    }.items():
        try:
            hs = word.styles[f"Heading {level_num}"]
            hs.font.name = "Inter"
            hs.font.size = Pt(size)
            hs.font.color.rgb = RGBColor(
                int(color_hex[0:2], 16),
                int(color_hex[2:4], 16),
                int(color_hex[4:6], 16),
            )
            hs.font.bold = True
            hs.paragraph_format.space_before = Pt(16 if level_num <= 1 else 12)
            hs.paragraph_format.space_after = Pt(6)
        except KeyError:
            pass

    # List styles
    for list_style_name in ["List Bullet", "List Number"]:
        try:
            ls = word.styles[list_style_name]
            ls.font.name = "Inter"
            ls.font.size = Pt(10.5)
            ls.font.color.rgb = RGBColor(55, 65, 81)
        except KeyError:
            pass

    # Document title
    title_para = word.add_heading(doc.title or "Untitled Document", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Project name
    session = db.query(SessionModel).filter(
        SessionModel.id == doc.session_id
    ).first()
    project_name = session.document_filename or "Untitled" if session else ""

    # Metadata line
    meta_parts = []
    if project_name:
        meta_parts.append(f"Project: {project_name}")
    meta_parts.append(f"Status: {doc.status.value.title() if doc.status else 'Draft'}")
    if doc.created_at:
        meta_parts.append(f"Created: {doc.created_at.strftime('%B %d, %Y')}")
    meta_para = word.add_paragraph("  |  ".join(meta_parts))
    meta_run = meta_para.runs[0] if meta_para.runs else meta_para.add_run("")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(107, 114, 128)

    # Divider
    word.add_paragraph("─" * 60).runs[0].font.color.rgb = RGBColor(229, 231, 235)

    # Render content blocks
    for block in content_blocks:
        _render_block_to_docx_v2(word, block)

    # Pricing table
    pricing_items = (
        db.query(PricingLineItem)
        .filter(PricingLineItem.document_id == doc.id)
        .order_by(PricingLineItem.order_index)
        .all()
    )

    if pricing_items:
        word.add_heading("Pricing", level=1)
        table = word.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Item", "Qty", "Unit Price", "Discount", "Tax", "Total"]
        for idx_h, header in enumerate(headers):
            cell = table.rows[0].cells[idx_h]
            cell.text = header
            # Bold header styling
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.name = "Inter"
            # Gray background
            shading = cell._element.get_or_add_tcPr()
            bg = shading.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "F9FAFB",
            })
            shading.append(bg)

        grand_total = 0.0
        for item in pricing_items:
            line_total = item.quantity * item.unit_price * (1 - item.discount_percent / 100) * (1 + item.tax_percent / 100)
            grand_total += line_total
            row = table.add_row()
            row.cells[0].text = item.name[:50]
            row.cells[1].text = f"{item.quantity:.1f}"
            row.cells[2].text = f"${item.unit_price:,.2f}"
            row.cells[3].text = f"{item.discount_percent:.0f}%"
            row.cells[4].text = f"{item.tax_percent:.0f}%"
            row.cells[5].text = f"${line_total:,.2f}"

        total_row = table.add_row()
        total_row.cells[0].text = "Total"
        for p in total_row.cells[0].paragraphs:
            for run in p.runs:
                run.bold = True
        total_row.cells[5].text = f"${grand_total:,.2f}"

    # Output
    buf = io.BytesIO()
    word.save(buf)
    buf.seek(0)

    filename = f"{doc.title or 'document'}_{datetime.utcnow().strftime('%Y%m%d')}.docx"
    filename = "".join(c if c.isalnum() or c in "._- " else "_" for c in filename)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _extract_cell_text_export(cell) -> str:
    """Extract plain text from a table cell in any storage format."""
    if cell is None:
        return ""
    if isinstance(cell, str):
        return cell
    if isinstance(cell, list):
        if cell and isinstance(cell[0], list):
            return _extract_cell_text_export(cell[0])
        parts = []
        for item in cell:
            if isinstance(item, dict):
                parts.append(item.get("text", ""))
            elif isinstance(item, str):
                parts.append(item)
        return "".join(parts)
    if isinstance(cell, dict):
        if "content" in cell:
            return _extract_cell_text_export(cell["content"])
        return cell.get("text", "")
    return str(cell)


def _render_block_to_docx_v2(word, block: dict) -> None:
    """Render a single content block to a python-docx Document with proper styling."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    block_type = block.get("type", "paragraph")
    props = block.get("props", {})
    content = block.get("content", [])

    # Extract inline text with styling info
    inline_items = []
    if isinstance(content, list):
        for item in content:
            if isinstance(item, dict):
                inline_items.append(item)
            elif isinstance(item, str):
                inline_items.append({"type": "text", "text": item, "styles": {}})

    plain_text = "".join(item.get("text", "") for item in inline_items)

    if block_type == "heading":
        level = min(props.get("level", 1), 4)
        para = word.add_heading("", level=level)
        # Add inline content with formatting
        for item in inline_items:
            run = para.add_run(item.get("text", ""))
            styles = item.get("styles", {})
            if styles.get("bold"):
                run.bold = True
            if styles.get("italic"):
                run.italic = True

    elif block_type == "bulletListItem":
        para = word.add_paragraph("", style="List Bullet")
        for item in inline_items:
            run = para.add_run(item.get("text", ""))
            styles = item.get("styles", {})
            run.font.name = "Inter"
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(55, 65, 81)
            if styles.get("bold"):
                run.bold = True
            if styles.get("italic"):
                run.italic = True
            if styles.get("code"):
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(79, 70, 229)

    elif block_type == "numberedListItem":
        para = word.add_paragraph("", style="List Number")
        for item in inline_items:
            run = para.add_run(item.get("text", ""))
            styles = item.get("styles", {})
            run.font.name = "Inter"
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(55, 65, 81)
            if styles.get("bold"):
                run.bold = True
            if styles.get("italic"):
                run.italic = True

    elif block_type == "codeBlock":
        lang = props.get("language", "")
        para = word.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        # Code block with background
        run = para.add_run(plain_text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 41, 59)
        # Add shading to paragraph
        pPr = para._element.get_or_add_pPr()
        shading = pPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "F1F5F9",
        })
        pPr.append(shading)

    elif block_type == "mermaid":
        code = props.get("code", plain_text)
        # Styled placeholder for mermaid diagram
        para = word.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(4)
        # Box header
        run = para.add_run("📊 Diagram")
        run.bold = True
        run.font.name = "Inter"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(59, 130, 246)

        # Description line
        desc_para = word.add_paragraph()
        desc_run = desc_para.add_run("(This diagram is rendered as an interactive chart in the web editor. Below is the diagram definition.)")
        desc_run.font.name = "Inter"
        desc_run.font.size = Pt(8)
        desc_run.font.italic = True
        desc_run.font.color.rgb = RGBColor(156, 163, 175)

        # Diagram code in styled block
        code_para = word.add_paragraph()
        code_para.paragraph_format.space_after = Pt(8)
        code_run = code_para.add_run(code)
        code_run.font.name = "Courier New"
        code_run.font.size = Pt(8)
        code_run.font.color.rgb = RGBColor(30, 41, 59)
        # Shading
        pPr = code_para._element.get_or_add_pPr()
        shading = pPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "EFF6FF",
        })
        pPr.append(shading)

    elif block_type == "table":
        # Render BlockNote table
        table_content = content
        rows_data = []
        if isinstance(table_content, dict):
            rows_data = table_content.get("rows", [])
        elif isinstance(table_content, list):
            rows_data = table_content

        if rows_data:
            num_cols = max(len(r.get("cells", [])) for r in rows_data) if rows_data else 1
            table = word.add_table(rows=0, cols=num_cols)
            table.style = "Table Grid"

            for ri, row in enumerate(rows_data):
                cells = row.get("cells", [])
                word_row = table.add_row()
                for ci in range(num_cols):
                    cell_data = cells[ci] if ci < len(cells) else None
                    cell_text = _extract_cell_text_export(cell_data)
                    word_row.cells[ci].text = cell_text

                    # Style cell text
                    for p in word_row.cells[ci].paragraphs:
                        for run in p.runs:
                            run.font.name = "Inter"
                            run.font.size = Pt(9)
                            if ri == 0:
                                run.bold = True
                                run.font.color.rgb = RGBColor(55, 65, 81)
                            else:
                                run.font.color.rgb = RGBColor(75, 85, 99)

                    # Header row background
                    if ri == 0:
                        tc_pr = word_row.cells[ci]._element.get_or_add_tcPr()
                        shading = tc_pr.makeelement(qn("w:shd"), {
                            qn("w:val"): "clear",
                            qn("w:color"): "auto",
                            qn("w:fill"): "F9FAFB",
                        })
                        tc_pr.append(shading)

    elif block_type in ("quote", "callout"):
        para = word.add_paragraph()
        para.paragraph_format.left_indent = Pt(18)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        # Left border via shading
        pPr = para._element.get_or_add_pPr()
        shading = pPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "FAFAF5",
        })
        pPr.append(shading)
        for item in inline_items:
            run = para.add_run(item.get("text", ""))
            run.font.name = "Inter"
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(75, 85, 99)

    else:
        # Default paragraph with inline styling
        if plain_text:
            para = word.add_paragraph()
            for item in inline_items:
                run = para.add_run(item.get("text", ""))
                styles = item.get("styles", {})
                run.font.name = "Inter"
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(55, 65, 81)
                if styles.get("bold"):
                    run.bold = True
                    run.font.color.rgb = RGBColor(17, 24, 39)
                if styles.get("italic"):
                    run.italic = True
                if styles.get("code"):
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(79, 70, 229)
                if styles.get("underline"):
                    run.underline = True
                if styles.get("strikethrough"):
                    run.font.strike = True

    # Recurse into children
    for child in block.get("children", []):
        if isinstance(child, dict):
            _render_block_to_docx_v2(word, child)


@router.get("/{document_id}/pdf")
def export_document_pdf(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Render document to PDF and return as a streaming download."""
    doc = _verify_document_access(db, document_id, current_user)

    content_blocks = doc.content or []

    try:
        from fpdf import FPDF
    except ImportError:
        raise HTTPException(
            500,
            "PDF generation requires the fpdf2 package. Install with: pip install fpdf2",
        )

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Use built-in fonts (DejaVu available in container, but fallback to Helvetica for safety)
    PAGE_W = pdf.w - pdf.l_margin - pdf.r_margin

    # ── Title ──
    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(17, 24, 39)
    pdf.multi_cell(0, 10, doc.title or "Untitled Document")
    pdf.ln(2)

    # ── Metadata line ──
    session = db.query(SessionModel).filter(
        SessionModel.id == doc.session_id
    ).first()
    project_name = session.document_filename or "Untitled" if session else ""

    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(107, 114, 128)
    meta_parts = []
    if project_name:
        meta_parts.append(f"Project: {project_name}")
    meta_parts.append(f"Status: {doc.status.value.title() if doc.status else 'Draft'}")
    if doc.created_at:
        meta_parts.append(f"Created: {doc.created_at.strftime('%B %d, %Y')}")
    pdf.cell(0, 5, "  |  ".join(meta_parts), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    # Divider line
    pdf.set_draw_color(229, 231, 235)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(6)

    # Reset text color
    pdf.set_text_color(55, 65, 81)

    # ── Render content blocks ──
    _pdf_list_counter = [0]  # mutable counter for numbered lists

    for block in content_blocks:
        _render_block_to_pdf_v2(pdf, block, PAGE_W, _pdf_list_counter)

    # ── Pricing table ──
    pricing_items = (
        db.query(PricingLineItem)
        .filter(PricingLineItem.document_id == doc.id)
        .order_by(PricingLineItem.order_index)
        .all()
    )

    if pricing_items:
        pdf.ln(8)
        pdf.set_font("Helvetica", "B", 16)
        pdf.set_text_color(17, 24, 39)
        pdf.cell(0, 10, "Pricing", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        col_widths = [55, 18, 28, 25, 20, 34]
        headers = ["Item", "Qty", "Unit Price", "Discount", "Tax", "Total"]

        # Header row
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_fill_color(249, 250, 251)
        pdf.set_text_color(55, 65, 81)
        for w, h in zip(col_widths, headers):
            pdf.cell(w, 7, h, border=1, fill=True)
        pdf.ln()

        # Data rows
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(75, 85, 99)
        grand_total = 0.0
        for item in pricing_items:
            line_total = item.quantity * item.unit_price * (1 - item.discount_percent / 100) * (1 + item.tax_percent / 100)
            grand_total += line_total

            name_display = item.name[:28] + ".." if len(item.name) > 30 else item.name
            pdf.cell(col_widths[0], 7, name_display, border=1)
            pdf.cell(col_widths[1], 7, f"{item.quantity:.1f}", border=1)
            pdf.cell(col_widths[2], 7, f"${item.unit_price:,.2f}", border=1)
            pdf.cell(col_widths[3], 7, f"{item.discount_percent:.0f}%", border=1)
            pdf.cell(col_widths[4], 7, f"{item.tax_percent:.0f}%", border=1)
            pdf.cell(col_widths[5], 7, f"${line_total:,.2f}", border=1)
            pdf.ln()

        # Grand total row
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_fill_color(249, 250, 251)
        pdf.set_text_color(17, 24, 39)
        total_label_width = sum(col_widths[:-1])
        pdf.cell(total_label_width, 8, "Grand Total", border=1, fill=True, align="R")
        pdf.cell(col_widths[-1], 8, f"${grand_total:,.2f}", border=1, fill=True)
        pdf.ln()

    # Output PDF bytes
    pdf_bytes = pdf.output()

    filename = f"{doc.title or 'document'}_{datetime.utcnow().strftime('%Y%m%d')}.pdf"
    filename = "".join(c if c.isalnum() or c in "._- " else "_" for c in filename)

    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _safe_text(text: str) -> str:
    """Sanitize text for fpdf2: replace chars that can't be encoded in latin-1."""
    replacements = {
        "\u2013": "-", "\u2014": "--", "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"', "\u2026": "...", "\u2022": "-",
        "\u2192": "->", "\u2190": "<-", "\u2194": "<->",
        "\u2500": "-", "\u2502": "|", "\u250c": "+", "\u2510": "+",
        "\u2514": "+", "\u2518": "+", "\u251c": "+", "\u2524": "+",
        "\u252c": "+", "\u2534": "+", "\u253c": "+",
        "\u2713": "v", "\u2717": "x", "\u00a0": " ",
    }
    for ch, repl in replacements.items():
        text = text.replace(ch, repl)
    # Final fallback: replace any remaining non-latin1 chars
    try:
        text.encode("latin-1")
    except UnicodeEncodeError:
        text = text.encode("latin-1", errors="replace").decode("latin-1")
    return text


def _render_block_to_pdf_v2(pdf: "FPDF", block: dict, page_w: float, list_counter: list) -> None:
    """Render a single content block to the PDF with proper styling."""
    block_type = block.get("type", "paragraph")
    props = block.get("props", {})
    content = block.get("content", [])

    # Extract inline items
    inline_items = []
    if isinstance(content, list):
        for item in content:
            if isinstance(item, dict):
                inline_items.append(item)
            elif isinstance(item, str):
                inline_items.append({"type": "text", "text": item, "styles": {}})

    plain_text = _safe_text("".join(item.get("text", "") for item in inline_items))

    if block_type == "heading":
        level = props.get("level", 1)
        sizes = {1: 18, 2: 15, 3: 13, 4: 11}
        pdf.ln(4 if level <= 2 else 3)
        pdf.set_font("Helvetica", "B", sizes.get(level, 11))
        pdf.set_text_color(17, 24, 39)
        pdf.multi_cell(0, sizes.get(level, 11) * 0.5, _safe_text(plain_text))
        pdf.ln(2)

    elif block_type == "bulletListItem":
        list_counter[0] = 0  # reset numbered counter
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(55, 65, 81)
        x = pdf.get_x()
        pdf.cell(6, 6, "", new_x="END")  # indent
        pdf.cell(4, 6, "-", new_x="END")
        _pdf_write_inline(pdf, inline_items, page_w - 12)
        pdf.ln(1)

    elif block_type == "numberedListItem":
        list_counter[0] += 1
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(55, 65, 81)
        pdf.cell(6, 6, "", new_x="END")  # indent
        pdf.cell(6, 6, f"{list_counter[0]}.", new_x="END")
        _pdf_write_inline(pdf, inline_items, page_w - 14)
        pdf.ln(1)

    elif block_type == "codeBlock":
        pdf.ln(3)
        pdf.set_fill_color(241, 245, 249)  # slate-100
        pdf.set_font("Courier", "", 8)
        pdf.set_text_color(30, 41, 59)  # slate-800
        # Split into lines and render
        for line in plain_text.split("\n"):
            pdf.cell(0, 5, "  " + _safe_text(line), fill=True, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

    elif block_type == "mermaid":
        code = _safe_text(props.get("code", plain_text))
        pdf.ln(4)
        # Header
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(59, 130, 246)
        pdf.cell(0, 6, "Diagram", new_x="LMARGIN", new_y="NEXT")
        # Description
        pdf.set_font("Helvetica", "I", 7)
        pdf.set_text_color(156, 163, 175)
        pdf.cell(0, 4, "(Interactive chart in web editor. Diagram definition below.)", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)
        # Code block
        pdf.set_fill_color(239, 246, 255)  # blue-50
        pdf.set_font("Courier", "", 7)
        pdf.set_text_color(30, 41, 59)
        for line in code.split("\n"):
            pdf.cell(0, 4, "  " + _safe_text(line), fill=True, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    elif block_type == "table":
        table_content = content
        rows_data = []
        if isinstance(table_content, dict):
            rows_data = table_content.get("rows", [])
        elif isinstance(table_content, list):
            rows_data = table_content

        if rows_data:
            pdf.ln(3)
            num_cols = max(len(r.get("cells", [])) for r in rows_data) if rows_data else 1
            col_w = min(page_w / num_cols, 50)  # cap column width

            for ri, row in enumerate(rows_data):
                cells = row.get("cells", [])
                if ri == 0:
                    pdf.set_font("Helvetica", "B", 8)
                    pdf.set_fill_color(249, 250, 251)
                    pdf.set_text_color(55, 65, 81)
                else:
                    pdf.set_font("Helvetica", "", 8)
                    pdf.set_text_color(75, 85, 99)

                for ci in range(num_cols):
                    cell_data = cells[ci] if ci < len(cells) else None
                    cell_text = _safe_text(_extract_cell_text_export(cell_data))
                    # Truncate long cell text
                    if len(cell_text) > 40:
                        cell_text = cell_text[:38] + ".."
                    pdf.cell(col_w, 6, cell_text, border=1, fill=(ri == 0))
                pdf.ln()
            pdf.ln(3)

    elif block_type in ("quote", "callout"):
        pdf.ln(2)
        pdf.set_fill_color(250, 250, 245)
        pdf.set_font("Helvetica", "I", 10)
        pdf.set_text_color(75, 85, 99)
        x = pdf.get_x()
        # Left indent
        pdf.set_x(x + 8)
        pdf.multi_cell(page_w - 16, 6, plain_text, fill=True)
        pdf.ln(2)

    else:
        # Default paragraph with inline styling
        if plain_text:
            list_counter[0] = 0  # reset numbered counter
            _pdf_write_inline(pdf, inline_items, page_w)
            pdf.ln(3)

    # Reset default text color
    pdf.set_text_color(55, 65, 81)

    # Recurse into children
    for child in block.get("children", []):
        if isinstance(child, dict):
            _render_block_to_pdf_v2(pdf, child, page_w, list_counter)


def _pdf_write_inline(pdf: "FPDF", inline_items: list, max_w: float) -> None:
    """Write inline content items with bold/italic/code styling."""
    for item in inline_items:
        text = _safe_text(item.get("text", ""))
        if not text:
            continue
        styles = item.get("styles", {})

        is_bold = styles.get("bold", False)
        is_italic = styles.get("italic", False)
        is_code = styles.get("code", False)

        if is_code:
            pdf.set_font("Courier", "", 8)
            pdf.set_text_color(79, 70, 229)
        else:
            style_str = ""
            if is_bold:
                style_str += "B"
            if is_italic:
                style_str += "I"
            pdf.set_font("Helvetica", style_str, 10)

            if is_bold:
                pdf.set_text_color(17, 24, 39)
            else:
                pdf.set_text_color(55, 65, 81)

        pdf.write(6, text)

    # Reset
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(55, 65, 81)
    pdf.ln()


# -- AI Document Generation Endpoints ----------------------------------------

class AIMessageRequest(BaseModel):
    message: str = Field(..., min_length=1, max_length=5000)


class AIApplyRequest(BaseModel):
    blocks: list = Field(..., description="BlockNote JSON blocks to merge into the document")


@router.post("/{document_id}/ai/chat")
async def send_ai_chat_message(
    document_id: str,
    body: AIMessageRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Send a message to the AI assistant and auto-apply generated content to the document."""
    from app.services.document_ai_service import generate_document_content, apply_blocks_to_document

    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Verify session membership
    session = db.query(SessionModel).filter(SessionModel.id == doc.session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Project not found")

    try:
        result = await generate_document_content(
            db=db,
            document_id=UUID(document_id),
            user_message=body.message,
            user=current_user,
        )

        # Auto-apply generated blocks to the document
        updated_content = None
        has_blocks = result.get("blocks") and len(result["blocks"]) > 0
        is_full_replace = result.get("is_full_replace", False)

        if has_blocks:
            try:
                apply_result = await apply_blocks_to_document(
                    db=db,
                    document_id=UUID(document_id),
                    new_blocks=result.get("blocks", []),
                    user=current_user,
                    full_replace=is_full_replace,
                )
                updated_content = apply_result.get("content")
            except Exception as apply_err:
                logger.warning(f"Auto-apply failed, blocks still available: {apply_err}")

        result["auto_applied"] = updated_content is not None
        result["updated_content"] = updated_content
        return result
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.get("/{document_id}/ai/history")
def get_ai_chat_history(
    document_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Get AI chat history for a document."""
    from app.services.document_ai_service import get_chat_history

    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    return {"messages": get_chat_history(db, UUID(document_id))}


@router.post("/{document_id}/ai/apply")
async def apply_ai_blocks_to_document(
    document_id: str,
    body: AIApplyRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Intelligently merge AI-generated blocks into the document.

    Instead of appending at the end, this analyses the document structure
    and places new content where it logically belongs — replacing sections
    with matching headings and inserting new sections at appropriate positions.
    """
    from app.services.document_ai_service import apply_blocks_to_document

    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    try:
        result = await apply_blocks_to_document(
            db=db,
            document_id=UUID(document_id),
            new_blocks=body.blocks,
            user=current_user,
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.delete("/{document_id}/ai/history")
def clear_ai_chat_history(
    document_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Clear AI chat history for a document."""
    from app.services.document_ai_service import clear_chat_history

    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    clear_chat_history(db, UUID(document_id))
    return {"status": "cleared"}


@router.get("/{document_id}/preview")
def get_document_preview(
    document_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Get a rich HTML preview of the document, including Mermaid diagrams."""
    from app.services.document_ai_service import generate_document_preview_html
    from fastapi.responses import HTMLResponse

    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Verify session membership
    session = db.query(SessionModel).filter(SessionModel.id == doc.session_id).first()
    if not session:
        raise HTTPException(status_code=404, detail="Project not found")

    try:
        html = generate_document_preview_html(db, UUID(document_id))
        return HTMLResponse(content=html)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
