"""
API routes for source ingestion and suggestion management.
Phase 1: Meeting Notes Integration.

Supports:
- Text paste (JSON body)
- File upload (multipart form) — .txt, .md, .pdf, .docx, .json, .csv, .html
- Multiple source formats (Otter.ai, Fireflies.ai, email, Slack, raw)
- Meeting type metadata for better AI context
"""
import os
import io
from uuid import UUID
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Query
from sqlalchemy.orm import Session
from app.db.session import get_db
from app.core.auth import get_optional_user
from app.models.database import (
    Source, SourceSuggestion, SourceType, User
)
from app.models.schemas import (
    SourceIngestRequest, SourceResponse, SourceDetailResponse,
    SuggestionResponse, SuggestionApplyRequest, BulkSuggestionApplyRequest,
    SuccessResponse,
)
from app.services.meeting_notes_parser import meeting_notes_parser
from app.services.graph_service import graph_service
from app.core.logger import get_logger

logger = get_logger(__name__)

router = APIRouter(prefix="/sources", tags=["sources"])


# ===== Helper: Build Source Response =====

def _build_source_detail_response(source: Source, suggestions: list) -> dict:
    """Build a consistent source detail response dictionary."""
    suggestion_responses = []
    for suggestion in suggestions:
        suggestion_responses.append({
            "id": str(suggestion.id),
            "source_id": str(suggestion.source_id),
            "change_type": suggestion.change_type,
            "target_node_id": str(suggestion.target_node_id) if suggestion.target_node_id else None,
            "parent_node_id": str(suggestion.parent_node_id) if suggestion.parent_node_id else None,
            "title": suggestion.title,
            "description": suggestion.description,
            "acceptance_criteria": suggestion.acceptance_criteria or [],
            "confidence": suggestion.confidence,
            "reasoning": suggestion.reasoning,
            "source_quote": suggestion.source_quote,
            "is_approved": suggestion.is_approved,
            "reviewer_note": suggestion.reviewer_note,
        })

    return {
        "id": str(source.id),
        "session_id": str(source.session_id),
        "source_type": source.source_type.value,
        "source_name": source.source_name,
        "raw_content": source.raw_content,
        "processed_summary": source.processed_summary,
        "is_processed": source.is_processed,
        "source_metadata": source.source_metadata,
        "suggestions": suggestion_responses,
        "created_at": source.created_at,
    }


def _build_source_list_item(source: Source, suggestions: list) -> dict:
    """Build a consistent source list item response dictionary."""
    approved_count = sum(1 for s in suggestions if s.is_approved is True)
    rejected_count = sum(1 for s in suggestions if s.is_approved is False)
    pending_count = sum(1 for s in suggestions if s.is_approved is None)

    return {
        "id": str(source.id),
        "session_id": str(source.session_id),
        "source_type": source.source_type.value,
        "source_name": source.source_name,
        "is_processed": source.is_processed,
        "processed_summary": source.processed_summary,
        "source_metadata": source.source_metadata,
        "suggestions_count": len(suggestions),
        "approved_count": approved_count,
        "rejected_count": rejected_count,
        "pending_count": pending_count,
        "created_at": source.created_at,
    }


# ===== File Text Extraction =====

def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text content from a PDF file.
    Uses PyPDF2 to read all pages and concatenate their text.
    """
    try:
        from PyPDF2 import PdfReader

        pdf_reader = PdfReader(io.BytesIO(file_bytes))
        all_pages_text = []

        for page_number, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                all_pages_text.append(page_text.strip())

        extracted_text = '\n\n'.join(all_pages_text)
        logger.info(f"Extracted {len(extracted_text)} chars from PDF ({len(pdf_reader.pages)} pages)")
        return extracted_text

    except Exception as error:
        logger.error(f"Error extracting text from PDF: {error}")
        raise HTTPException(
            status_code=400,
            detail=f"Could not extract text from PDF: {str(error)}",
        )


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a Word .docx file (ZIP-based format).
    Uses python-docx to read all paragraphs.
    """
    try:
        from docx import Document

        document = Document(io.BytesIO(file_bytes))
        all_paragraphs_text = []

        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text.strip()
            if paragraph_text:
                all_paragraphs_text.append(paragraph_text)

        # Also extract text from tables (common in meeting notes docs)
        for table in document.tables:
            for row in table.rows:
                row_cells_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells_text.append(cell_text)
                if row_cells_text:
                    all_paragraphs_text.append(' | '.join(row_cells_text))

        extracted_text = '\n\n'.join(all_paragraphs_text)
        logger.info(f"Extracted {len(extracted_text)} chars from DOCX ({len(all_paragraphs_text)} paragraphs)")
        return extracted_text

    except Exception as error:
        logger.error(f"Error extracting text from DOCX: {error}")
        raise HTTPException(
            status_code=400,
            detail=f"Could not extract text from .docx file: {str(error)}",
        )


def _extract_text_from_doc_impl(file_bytes: bytes) -> str:
    """
    Extract text from an old-format Word .doc file (binary OLE format).

    python-docx does NOT support .doc. We try:
    1. UTF-16-LE (common in Office .doc)
    2. Latin-1 / raw byte scan for ASCII runs
    3. Alternate byte layout (low bytes of UTF-16)
    4. UTF-16-BE
    """
    import re

    def _collect_runs(text: str, min_run: int) -> list:
        """Extract runs of printable chars (letters, digits, punctuation, space)."""
        runs = re.findall(r'[\x20-\x7E\n\r\t]{' + str(min_run) + r',}', text)
        return runs

    def _is_noise(s: str) -> bool:
        """Filter out likely binary/formatting noise."""
        if len(s) < 3:
            return True
        # All same character (e.g. --------)
        non_ws = s.replace(' ', '').replace('\n', '').replace('\t', '')
        if non_ws and len(set(non_ws)) <= 1:
            return True
        # At least some letters or digits (avoid pure binary)
        letters_or_digits = sum(1 for c in s if c.isalnum() or c.isspace())
        if letters_or_digits < len(s) * 0.2:
            return True
        return False

    all_runs = []

    # 1. Try UTF-16 with BOM (Word sometimes writes BOM)
    try:
        utf16_text = file_bytes.decode('utf-16', errors='ignore')
        runs = _collect_runs(utf16_text, min_run=3)
        all_runs.extend(r for r in runs if not _is_noise(r))
    except Exception:
        pass

    # 2. Try UTF-16-LE without BOM (common for .doc)
    try:
        utf16_text = file_bytes.decode('utf-16-le', errors='ignore')
        runs = _collect_runs(utf16_text, min_run=3)
        all_runs.extend(r for r in runs if not _is_noise(r))
    except Exception:
        pass

    # 3. Latin-1 decode and extract ASCII runs (short and long)
    try:
        raw_text = file_bytes.decode('latin-1', errors='replace')
        for min_len in (4, 8, 15, 20):
            runs = _collect_runs(raw_text, min_run=min_len)
            for r in runs:
                if not _is_noise(r) and r not in all_runs:
                    all_runs.append(r)
    except Exception:
        pass

    # 4. Low bytes of UTF-16 LE (every other byte) — some .doc store ASCII as 16-bit
    if len(file_bytes) >= 2:
        try:
            low_bytes = bytes(file_bytes[i] for i in range(0, len(file_bytes), 2))
            decoded = low_bytes.decode('latin-1', errors='replace')
            runs = _collect_runs(decoded, min_run=4)
            all_runs.extend(r for r in runs if not _is_noise(r) and r not in all_runs)
        except Exception:
            pass

    # 5. UTF-16-BE
    try:
        utf16be_text = file_bytes.decode('utf-16-be', errors='ignore')
        runs = _collect_runs(utf16be_text, min_run=3)
        for r in runs:
            if not _is_noise(r) and r not in all_runs:
                all_runs.append(r)
    except Exception:
        pass

    # Deduplicate while preserving order (keep first occurrence)
    seen = set()
    unique_runs = []
    for r in all_runs:
        r_stripped = r.strip()
        if r_stripped and r_stripped not in seen:
            seen.add(r_stripped)
            unique_runs.append(r_stripped)

    if not unique_runs:
        raise HTTPException(
            status_code=400,
            detail="Could not extract readable text from .doc file. "
                   "Please save it as .docx or export as PDF and try again.",
        )

    extracted_text = '\n\n'.join(unique_runs)
    extracted_text = re.sub(r'\n{3,}', '\n\n', extracted_text).strip()

    logger.info(f"Extracted {len(extracted_text)} chars from .doc ({len(unique_runs)} segments)")
    return extracted_text


def _extract_text_from_doc(file_bytes: bytes) -> str:
    """Wrapper to catch any unexpected errors during .doc extraction."""
    try:
        return _extract_text_from_doc_impl(file_bytes)
    except HTTPException:
        raise
    except Exception as error:
        logger.error(f"Error extracting text from .doc: {error}")
        raise HTTPException(
            status_code=400,
            detail="Could not extract text from .doc file. "
                   "Please save as .docx or PDF and try again.",
        )


def _extract_text_from_word(file_bytes: bytes, file_extension: str) -> str:
    """
    Smart Word file extractor — tries .docx first, falls back to .doc.
    Handles the common case where a .doc file is renamed to .docx or vice-versa.
    """
    if file_extension == '.doc':
        # Old format — use binary extraction
        return _extract_text_from_doc(file_bytes)

    # .docx — try python-docx first
    try:
        return _extract_text_from_docx(file_bytes)
    except HTTPException:
        # If python-docx fails (e.g., "not a zip file"), the file is
        # probably an old .doc saved with a .docx extension. Fall back.
        logger.warning(
            "python-docx failed on .docx file — falling back to .doc binary extraction. "
            "The file is likely an old Word .doc format with a .docx extension."
        )
        return _extract_text_from_doc(file_bytes)


# ===== Ingestion Endpoints =====

@router.post("/ingest", response_model=SourceDetailResponse)
async def ingest_source(
    request: SourceIngestRequest,
    database: Session = Depends(get_db),
    current_user: User = Depends(get_optional_user),
):
    """
    Ingest a source via text paste (JSON body).
    AI analyzes the content and generates scope change suggestions.
    """
    try:
        # Build metadata including new fields
        source_metadata = request.source_metadata or {}
        source_metadata['source_format'] = request.source_format
        if request.meeting_type:
            source_metadata['meeting_type'] = request.meeting_type

        # Map source_type string to enum (with fallback)
        try:
            source_type_enum = SourceType(request.source_type)
        except ValueError:
            source_type_enum = SourceType.MEETING

        # Create source record
        source = Source(
            session_id=request.session_id,
            source_type=source_type_enum,
            source_name=request.source_name,
            raw_content=request.raw_content,
            source_metadata=source_metadata,
            created_by=current_user.id if current_user else None,
        )
        database.add(source)
        database.commit()
        database.refresh(source)

        logger.info(
            f"Created source '{source.source_name}' "
            f"(type={request.source_type}, format={request.source_format}) "
            f"for session {request.session_id}"
        )

        # Run AI analysis
        suggestions = await meeting_notes_parser.analyze_source(source.id, database)

        return _build_source_detail_response(source, suggestions)

    except Exception as error:
        logger.error(f"Error ingesting source: {str(error)}")
        raise HTTPException(status_code=500, detail=str(error))


@router.post("/upload", response_model=SourceDetailResponse)
async def upload_source_file(
    session_id: str = Form(...),
    source_name: str = Form(...),
    source_type: str = Form("meeting"),
    source_format: str = Form("raw"),
    meeting_type: str = Form(None),
    meeting_date: str = Form(None),
    attendees: str = Form(None),
    file: UploadFile = File(...),
    database: Session = Depends(get_db),
    current_user: User = Depends(get_optional_user),
):
    """
    Ingest a source via file upload (multipart form).
    Supports .txt, .md, .pdf, .docx, .json, .csv, .html files.
    AI analyzes the content and generates scope change suggestions.
    """
    try:
        # Validate file type
        text_extensions = {'.txt', '.md', '.json', '.csv', '.html', '.text', '.log'}
        binary_extensions = {'.pdf', '.docx', '.doc'}
        allowed_extensions = text_extensions | binary_extensions
        file_extension = os.path.splitext(file.filename or '')[1].lower()

        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type '{file_extension}'. "
                       f"Allowed: {', '.join(sorted(allowed_extensions))}",
            )

        # Read file bytes
        file_bytes = await file.read()

        # Extract text based on file type
        if file_extension == '.pdf':
            raw_content = _extract_text_from_pdf(file_bytes)
        elif file_extension in ('.docx', '.doc'):
            raw_content = _extract_text_from_word(file_bytes, file_extension)
        else:
            raw_content = file_bytes.decode('utf-8', errors='replace')

        if not raw_content.strip():
            raise HTTPException(status_code=400, detail="Uploaded file is empty or text could not be extracted.")

        # Build metadata
        source_metadata: dict = {
            'source_format': source_format,
            'uploaded_filename': file.filename,
            'file_size_bytes': len(file_bytes),
        }
        if meeting_type:
            source_metadata['meeting_type'] = meeting_type
        if meeting_date:
            source_metadata['date'] = meeting_date
        if attendees:
            source_metadata['attendees'] = attendees

        # Map source_type string to enum
        try:
            source_type_enum = SourceType(source_type)
        except ValueError:
            source_type_enum = SourceType.MEETING

        # Create source record
        source = Source(
            session_id=session_id,
            source_type=source_type_enum,
            source_name=source_name,
            raw_content=raw_content,
            source_metadata=source_metadata,
            created_by=current_user.id if current_user else None,
        )
        database.add(source)
        database.commit()
        database.refresh(source)

        logger.info(
            f"Uploaded source '{source.source_name}' from file '{file.filename}' "
            f"({len(file_bytes)} bytes) for session {session_id}"
        )

        # Run AI analysis
        suggestions = await meeting_notes_parser.analyze_source(source.id, database)

        return _build_source_detail_response(source, suggestions)

    except HTTPException:
        raise
    except Exception as error:
        logger.error(f"Error uploading source file: {str(error)}")
        raise HTTPException(status_code=500, detail=str(error))


# ===== Source Listing & Detail =====

@router.get("/{session_id}", response_model=None)
async def list_sources(
    session_id: str,
    database: Session = Depends(get_db),
    limit: int = Query(50, ge=1, le=200, description="Max items per page"),
    offset: int = Query(0, ge=0, description="Number of items to skip"),
):
    """List all sources for a session with suggestion counts (excludes soft-deleted).
    Supports PERF-3.1 pagination via ``limit`` and ``offset``.
    """
    try:
        base_query = (
            database.query(Source)
            .filter(Source.session_id == session_id)
            .order_by(Source.created_at.desc())
        )

        # Total count before filtering soft-deleted (conservative upper bound)
        all_sources = base_query.all()

        # Filter out soft-deleted sources
        active_sources = [
            source for source in all_sources
            if not (source.source_metadata or {}).get("is_deleted", False)
        ]

        total = len(active_sources)
        page = active_sources[offset : offset + limit]

        result = []
        for source in page:
            suggestions = (
                database.query(SourceSuggestion)
                .filter(SourceSuggestion.source_id == source.id)
                .all()
            )
            result.append(_build_source_list_item(source, suggestions))

        return {"items": result, "total": total, "limit": limit, "offset": offset}

    except Exception as error:
        logger.error(f"Error listing sources: {str(error)}")
        raise HTTPException(status_code=500, detail=str(error))


@router.get("/detail/{source_id}")
async def get_source_detail(
    source_id: str,
    database: Session = Depends(get_db),
):
    """Get a source with all its suggestions."""
    try:
        source = database.query(Source).filter(Source.id == source_id).first()
        if not source:
            raise HTTPException(status_code=404, detail="Source not found")

        suggestions = (
            database.query(SourceSuggestion)
            .filter(SourceSuggestion.source_id == source.id)
            .order_by(SourceSuggestion.confidence.desc())
            .all()
        )

        return _build_source_detail_response(source, suggestions)

    except HTTPException:
        raise
    except Exception as error:
        logger.error(f"Error getting source detail: {str(error)}")
        raise HTTPException(status_code=500, detail=str(error))


@router.get("/{source_id}/suggestions")
async def get_source_suggestions(
    source_id: str,
    database: Session = Depends(get_db),
):
    """Get all AI-generated suggestions for a specific source."""
    try:
        source = database.query(Source).filter(Source.id == source_id).first()
        if not source:
            raise HTTPException(status_code=404, detail="Source not found")

        suggestions = (
            database.query(SourceSuggestion)
            .filter(SourceSuggestion.source_id == source.id)
            .order_by(SourceSuggestion.confidence.desc())
            .all()
        )

        return [
            {
                "id": str(suggestion.id),
                "source_id": str(suggestion.source_id),
                "change_type": suggestion.change_type,
                "target_node_id": str(suggestion.target_node_id) if suggestion.target_node_id else None,
                "parent_node_id": str(suggestion.parent_node_id) if suggestion.parent_node_id else None,
                "title": suggestion.title,
                "description": suggestion.description,
                "acceptance_criteria": suggestion.acceptance_criteria or [],
                "confidence": suggestion.confidence,
                "reasoning": suggestion.reasoning,
                "source_quote": suggestion.source_quote,
                "is_approved": suggestion.is_approved,
            }
            for suggestion in suggestions
        ]

    except HTTPException:
        raise
    except Exception as error:
        logger.error(f"Error getting suggestions: {str(error)}")
        raise HTTPException(status_code=500, detail=str(error))


# ===== Source Deletion =====

@router.delete("/{source_id}")
async def delete_source(
    source_id: str,
    database: Session = Depends(get_db),
):
    """Soft-delete a source by marking it and clearing pending suggestions."""
    try:
        source = database.query(Source).filter(Source.id == source_id).first()
        if not source:
            raise HTTPException(status_code=404, detail="Source not found")

        # Delete all pending (unapproved) suggestions for this source
        database.query(SourceSuggestion).filter(
            SourceSuggestion.source_id == source.id,
            SourceSuggestion.is_approved.is_(None),
        ).delete(synchronize_session="fetch")

        # Soft-delete: mark source as inactive via metadata flag
        existing_metadata = source.source_metadata or {}
        existing_metadata["is_deleted"] = True
        existing_metadata["deleted_at"] = str(datetime.utcnow())
        source.source_metadata = existing_metadata

        database.commit()

        logger.info(f"Soft-deleted source {source_id}")
        return {"success": True, "message": f"Source '{source.source_name}' deleted"}

    except HTTPException:
        raise
    except Exception as error:
        logger.error(f"Error deleting source: {str(error)}")
        raise HTTPException(status_code=500, detail=str(error))


# ===== Re-analyze Source =====

@router.post("/{source_id}/reanalyze", response_model=SourceDetailResponse)
async def reanalyze_source(
    source_id: str,
    database: Session = Depends(get_db),
    current_user: User = Depends(get_optional_user),
):
    """
    Re-run AI analysis on an existing source.
    Clears all pending (un-reviewed) suggestions, then generates fresh ones.
    Already-approved/rejected suggestions are preserved for audit trail.
    """
    try:
        source = database.query(Source).filter(Source.id == source_id).first()
        if not source:
            raise HTTPException(status_code=404, detail="Source not found")

        # Refuse if the source was soft-deleted
        if (source.source_metadata or {}).get("is_deleted", False):
            raise HTTPException(status_code=400, detail="Cannot re-analyze a deleted source")

        # Delete only pending (un-reviewed) suggestions
        deleted_count = (
            database.query(SourceSuggestion)
            .filter(
                SourceSuggestion.source_id == source.id,
                SourceSuggestion.is_approved.is_(None),
            )
            .delete(synchronize_session="fetch")
        )
        database.commit()

        logger.info(
            f"Re-analyzing source '{source.source_name}' — "
            f"cleared {deleted_count} pending suggestions"
        )

        # Re-run AI analysis (creates new suggestion rows)
        new_suggestions = await meeting_notes_parser.analyze_source(source.id, database)

        # Return full detail including both old reviewed + new pending suggestions
        all_suggestions = (
            database.query(SourceSuggestion)
            .filter(SourceSuggestion.source_id == source.id)
            .order_by(SourceSuggestion.confidence.desc())
            .all()
        )

        return _build_source_detail_response(source, all_suggestions)

    except HTTPException:
        raise
    except Exception as error:
        logger.error(f"Error re-analyzing source: {str(error)}")
        raise HTTPException(status_code=500, detail=str(error))


# ===== Suggestion Application =====

@router.post("/apply", response_model=SuccessResponse)
async def apply_suggestions(
    request: BulkSuggestionApplyRequest,
    database: Session = Depends(get_db),
    current_user: User = Depends(get_optional_user),
):
    """
    Apply (approve/reject) a batch of suggestions.
    Approved suggestions modify the graph with full source attribution.
    """
    try:
        applied_nodes = []
        rejected_count = 0

        for decision in request.decisions:
            result = graph_service.apply_suggestion(
                database=database,
                suggestion_id=decision.suggestion_id,
                is_approved=decision.is_approved,
                approved_by=current_user.id if current_user else None,
                edited_title=decision.edited_title,
                edited_description=decision.edited_description,
                reviewer_note=decision.reviewer_note,
            )

            if result:
                applied_nodes.append(result)
            elif not decision.is_approved:
                rejected_count += 1

        message = f"Applied {len(applied_nodes)} changes, rejected {rejected_count} suggestions"
        logger.info(message)

        return SuccessResponse(
            success=True,
            message=message,
            data={
                "applied_count": len(applied_nodes),
                "rejected_count": rejected_count,
                "applied_nodes": applied_nodes,
            },
        )

    except Exception as error:
        logger.error(f"Error applying suggestions: {str(error)}")
        raise HTTPException(status_code=500, detail=str(error))
