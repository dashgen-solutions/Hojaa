"""
SVC-2.3 — Source Ingestion Service.

Coordinates the full ingestion pipeline:
1. Create Source record (text or file bytes)
2. Extract text from binary files (PDF, DOCX, DOC)
3. Pre-process with format-specific parser
4. Trigger scope-change detection via ScopeChangeDetector
5. Store suggestions and metadata
6. Return unified response

Extracts the orchestration logic that previously lived in the
route handlers (sources.py) into a testable, reusable service.
"""
import io
import os
from typing import Dict, Any, Optional, List
from uuid import UUID
from datetime import datetime
from sqlalchemy.orm import Session as DBSession
from app.models.database import (
    Source, SourceSuggestion, SourceType
)
from app.services.meeting_notes_parser import meeting_notes_parser
from app.core.logger import get_logger

logger = get_logger(__name__)


class SourceIngestionService:
    """
    Coordinates parsing, analysis, and storage of ingested sources.
    """

    # ── Text extraction ─────────────────────────────────────────

    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Extract text from PDF bytes using PyPDF2."""
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = [p.extract_text() or "" for p in reader.pages]
        text = "\n\n".join(t.strip() for t in pages if t.strip())
        logger.info(f"PDF: extracted {len(text)} chars from {len(reader.pages)} pages")
        return text

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Extract text from DOCX bytes using python-docx."""
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n\n".join(parts)
        logger.info(f"DOCX: extracted {len(text)} chars")
        return text

    @staticmethod
    def extract_text_from_file(file_bytes: bytes, extension: str) -> str:
        """
        Dispatch text extraction based on file extension.

        Raises ValueError if the extension is unsupported.
        """
        ext = extension.lower()
        svc = SourceIngestionService
        if ext == ".pdf":
            return svc.extract_text_from_pdf(file_bytes)
        if ext in (".docx",):
            return svc.extract_text_from_docx(file_bytes)
        if ext in (".txt", ".md", ".json", ".csv", ".html", ".text", ".log"):
            return file_bytes.decode("utf-8", errors="replace")
        raise ValueError(f"Unsupported file extension: {ext}")

    # ── Ingestion pipeline ──────────────────────────────────────

    async def ingest_text(
        self,
        database: DBSession,
        session_id: UUID,
        source_name: str,
        raw_content: str,
        source_type: str = "meeting",
        source_format: str = "raw",
        meeting_type: Optional[str] = None,
        source_metadata: Optional[Dict[str, Any]] = None,
        created_by: Optional[UUID] = None,
    ) -> Dict[str, Any]:
        """
        Ingest plain-text content.

        Returns:
            {source: Source, suggestions: [SourceSuggestion]}
        """
        metadata = source_metadata or {}
        metadata["source_format"] = source_format
        if meeting_type:
            metadata["meeting_type"] = meeting_type

        try:
            source_type_enum = SourceType(source_type)
        except ValueError:
            source_type_enum = SourceType.MEETING

        source = Source(
            session_id=session_id,
            source_type=source_type_enum,
            source_name=source_name,
            raw_content=raw_content,
            source_metadata=metadata,
            created_by=created_by,
        )
        database.add(source)
        database.commit()
        database.refresh(source)

        logger.info(
            f"Created source '{source_name}' "
            f"(type={source_type}, format={source_format}) "
            f"for session {session_id}"
        )

        suggestions = await self._run_analysis(source, database)
        return {"source": source, "suggestions": suggestions}

    async def ingest_file(
        self,
        database: DBSession,
        session_id: UUID,
        source_name: str,
        file_bytes: bytes,
        filename: str,
        source_type: str = "meeting",
        source_format: str = "raw",
        meeting_type: Optional[str] = None,
        meeting_date: Optional[str] = None,
        attendees: Optional[str] = None,
        created_by: Optional[UUID] = None,
    ) -> Dict[str, Any]:
        """
        Ingest a file upload: extract text, create Source, run analysis.

        Returns:
            {source: Source, suggestions: [SourceSuggestion]}
        """
        ext = os.path.splitext(filename or "")[1].lower()
        raw_content = self.extract_text_from_file(file_bytes, ext)

        if not raw_content.strip():
            raise ValueError("Uploaded file is empty or text could not be extracted.")

        metadata: Dict[str, Any] = {
            "source_format": source_format,
            "uploaded_filename": filename,
            "file_size_bytes": len(file_bytes),
        }
        if meeting_type:
            metadata["meeting_type"] = meeting_type
        if meeting_date:
            metadata["date"] = meeting_date
        if attendees:
            metadata["attendees"] = attendees

        try:
            source_type_enum = SourceType(source_type)
        except ValueError:
            source_type_enum = SourceType.MEETING

        source = Source(
            session_id=session_id,
            source_type=source_type_enum,
            source_name=source_name,
            raw_content=raw_content,
            source_metadata=metadata,
            created_by=created_by,
        )
        database.add(source)
        database.commit()
        database.refresh(source)

        logger.info(
            f"Ingested file '{filename}' as source '{source_name}' "
            f"({len(file_bytes)} bytes) for session {session_id}"
        )

        suggestions = await self._run_analysis(source, database)
        return {"source": source, "suggestions": suggestions}

    # ── Helpers ──────────────────────────────────────────────────

    async def _run_analysis(
        self,
        source: Source,
        database: DBSession,
    ) -> List[SourceSuggestion]:
        """Delegate analysis to MeetingNotesParser (which uses ScopeChangeDetector internally)."""
        return await meeting_notes_parser.analyze_source(source.id, database)

    def get_source_with_suggestions(
        self,
        database: DBSession,
        source_id: UUID,
    ) -> Optional[Dict[str, Any]]:
        """
        Load a source and its suggestions.

        Returns None if not found.
        """
        source = database.query(Source).filter(Source.id == source_id).first()
        if not source:
            return None
        suggestions = (
            database.query(SourceSuggestion)
            .filter(SourceSuggestion.source_id == source_id)
            .all()
        )
        return {"source": source, "suggestions": suggestions}

    def list_sources(
        self,
        database: DBSession,
        session_id: UUID,
        include_deleted: bool = False,
    ) -> List[Source]:
        """Return active sources for a session, optionally including soft-deleted."""
        query = (
            database.query(Source)
            .filter(Source.session_id == session_id)
            .order_by(Source.created_at.desc())
        )
        sources = query.all()
        if include_deleted:
            return sources
        return [
            s for s in sources
            if not (s.source_metadata or {}).get("is_deleted", False)
        ]

    def soft_delete_source(
        self,
        database: DBSession,
        source_id: UUID,
    ) -> bool:
        """Soft-delete a source by setting metadata flag. Returns True if found."""
        source = database.query(Source).filter(Source.id == source_id).first()
        if not source:
            return False
        meta = source.source_metadata or {}
        meta["is_deleted"] = True
        meta["deleted_at"] = datetime.utcnow().isoformat()
        source.source_metadata = meta
        database.commit()
        return True


# Global instance
source_ingestion_service = SourceIngestionService()
